#!/usr/bin/env python3
"""
Multi-platform file exporter for Writer.

Exports articles to platform-specific formats (HTML/Markdown/Word) with images.

Usage:
    python exporter.py export article.md --platform wechat --output ./output/2026-05-09-article/
    python exporter.py export article.md --platform all --output ./output/2026-05-09-article/
"""

import sys
import json
from pathlib import Path
from io import BytesIO

import yaml

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

from converter import WeChatConverter, preview_html
from theme import load_theme

# Platform display names for file naming
PLATFORM_DISPLAY_NAMES = {
    "wechat": "公众号",
    "xiaohongshu": "小红书",
    "zhihu": "知乎",
    "baijiahao": "百家号",
    "weibo": "微博",
    "sohu": "搜狐",
    "toutiao": "今日头条",
    "qiehao": "企鹅号",
    "jianshu": "简书",
    "douban": "豆瓣",
    "dayu": "大鱼号",
    "kr36": "36氪",
    "bilibili": "B站",
    "douyin": "抖音",
    "newsletter": "Newsletter",
}

SUPPORTED_PLATFORMS = [
    "wechat", "xiaohongshu", "zhihu", "baijiahao",
    "weibo", "sohu", "toutiao", "qiehao",
    "jianshu", "douban", "dayu", "kr36",
]

PLATFORM_FORMATS = {
    "wechat": "html",
    "xiaohongshu": "md",
    "zhihu": "md",
    "baijiahao": "html",
    "weibo": "md",
    "sohu": "html",
    "toutiao": "html",
    "qiehao": "html",
    "jianshu": "md",
    "douban": "md",
    "dayu": "html",
    "kr36": "md",
}

CONFIG_PATHS = [
    Path.cwd() / "config.yaml",
    Path(__file__).parent.parent / "config.yaml",
    Path(__file__).parent / "config.yaml",
    Path.home() / ".config" / "writer" / "config.yaml",
]


def _load_config() -> dict:
    for p in CONFIG_PATHS:
        if p.exists():
            with open(p, "r", encoding="utf-8") as f:
                return yaml.safe_load(f) or {}
    return {}


def _load_style() -> dict:
    for p in [Path.cwd() / "style.yaml", Path(__file__).parent.parent / "style.yaml"]:
        if p.exists():
            with open(p, "r", encoding="utf-8") as f:
                return yaml.safe_load(f) or {}
    return {}


def _convert_to_html(md_content: str, theme_name: str = "professional-clean") -> str:
    """Convert Markdown to WeChat-compatible HTML with inline styles."""
    theme = load_theme(theme_name)
    converter = WeChatConverter(theme=theme)
    result = converter.convert(md_content)
    return preview_html(result.html, theme)


def _convert_to_standard_html(md_content: str) -> str:
    """Convert Markdown to standard HTML without WeChat-specific fixes."""
    import markdown
    html_body = markdown.markdown(
        md_content,
        extensions=["tables", "fenced_code", "codehilite", "toc"],
        output_format="html"
    )
    return f"""<!DOCTYPE html>
<html lang="zh-CN">
<head>
<meta charset="UTF-8">
<meta name="viewport" content="width=device-width, initial-scale=1.0">
<style>
body {{ font-family: -apple-system, BlinkMacSystemFont, "Segoe UI", "PingFang SC", "Hiragino Sans GB", "Microsoft YaHei", sans-serif; line-height: 1.8; color: #333; max-width: 800px; margin: 0 auto; padding: 20px; }}
h1 {{ font-size: 24px; margin-top: 2em; }}
h2 {{ font-size: 20px; margin-top: 1.5em; }}
h3 {{ font-size: 18px; margin-top: 1.2em; }}
p {{ margin: 1em 0; }}
img {{ max-width: 100%; height: auto; display: block; margin: 1em auto; }}
blockquote {{ border-left: 4px solid #ddd; padding-left: 16px; color: #666; margin: 1em 0; }}
code {{ background: #f5f5f5; padding: 2px 6px; border-radius: 3px; font-family: monospace; }}
pre {{ background: #f5f5f5; padding: 16px; border-radius: 6px; overflow-x: auto; }}
pre code {{ background: none; padding: 0; }}
table {{ border-collapse: collapse; width: 100%; margin: 1em 0; }}
th, td {{ border: 1px solid #ddd; padding: 8px 12px; text-align: left; }}
th {{ background: #f5f5f5; font-weight: 600; }}
a {{ color: #1a73e8; text-decoration: none; }}
a:hover {{ text-decoration: underline; }}
ul, ol {{ padding-left: 2em; }}
</style>
</head>
<body>
{html_body}
</body>
</html>"""


def _process_markdown_for_platform(md_content: str, platform: str) -> str:
    """Process Markdown content for specific platform requirements."""
    if platform == "xiaohongshu":
        processed = _optimize_for_xiaohongshu(md_content)
    elif platform == "weibo":
        processed = _optimize_for_weibo(md_content)
    elif platform == "jianshu":
        processed = md_content
    elif platform == "douban":
        processed = md_content
    elif platform == "zhihu":
        processed = md_content
    elif platform == "kr36":
        processed = md_content
    else:
        processed = md_content
    return processed


def _optimize_for_xiaohongshu(md_content: str) -> str:
    """Optimize content for Xiaohongshu: emoji, short paragraphs, tags."""
    import re

    lines = md_content.split("\n")
    result_lines = []
    for line in lines:
        if line.startswith("# "):
            title = line[2:]
            result_lines.append(f"# {title}")
            result_lines.append("")
        elif line.startswith("## "):
            heading = line[3:]
            emoji = _get_section_emoji(heading)
            result_lines.append(f"## {emoji} {heading}")
            result_lines.append("")
        elif line.strip() and not line.startswith("#") and not line.startswith("-") and not line.startswith(">"):
            if line.strip().startswith("!["):
                result_lines.append("")
                result_lines.append(line)
                result_lines.append("")
            else:
                result_lines.append(line)
                result_lines.append("")
        else:
            result_lines.append(line)

    return "\n".join(result_lines)


def _get_section_emoji(heading: str) -> str:
    emoji_map = {
        "总结": "✨", "前言": "📌", "推荐": "🌟", "测评": "📊",
        "对比": "⚖️", "教程": "📝", "经验": "💡", "避坑": "⚠️",
        "工具": "🛠️", "方法": "🔧", "技巧": "💫", "建议": "✅",
        "注意": "🔥", "重点": "❗", "福利": "🎁", "案例": "📋",
    }
    for keyword, emoji in emoji_map.items():
        if keyword in heading:
            return emoji
    return "📖"


def _optimize_for_weibo(md_content: str) -> str:
    """Optimize content for Weibo: add hashtags, keep it short."""
    import re

    first_line = md_content.split("\n")[0]
    if first_line.startswith("# "):
        title = first_line[2:]
        hashtag = f"#{title.replace(' ', '')}#"
        md_content = md_content.replace(first_line, f"{hashtag}\n\n{title}", 1)

    return md_content


def _generate_word_from_html(html_content: str, output_path: str):
    """Generate Word document from HTML content."""
    from bs4 import BeautifulSoup
    import markdown

    soup = BeautifulSoup(html_content, "html.parser")

    doc = Document()

    doc.styles["Normal"].font.name = "微软雅黑"
    doc.styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    doc.styles["Normal"].font.size = Pt(12)
    doc.styles["Normal"].paragraph_format.space_after = Pt(6)
    doc.styles["Normal"].paragraph_format.line_spacing = 1.5

    heading_styles = {
        "h1": {"size": Pt(22), "bold": True, "color": RGBColor(0x1A, 0x1A, 0x1A), "space_before": Pt(18), "space_after": Pt(12)},
        "h2": {"size": Pt(18), "bold": True, "color": RGBColor(0x33, 0x33, 0x33), "space_before": Pt(14), "space_after": Pt(8)},
        "h3": {"size": Pt(15), "bold": True, "color": RGBColor(0x4D, 0x4D, 0x4D), "space_before": Pt(12), "space_after": Pt(6)},
        "h4": {"size": Pt(14), "bold": True, "color": RGBColor(0x4D, 0x4D, 0x4D), "space_before": Pt(10), "space_after": Pt(6)},
    }

    for element in soup.find_all(True):
        tag = element.name.lower()

        if tag in heading_styles:
            style = heading_styles[tag]
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_before = style["space_before"]
            p.paragraph_format.space_after = style["space_after"]
            run = p.add_run(element.get_text())
            run.font.size = style["size"]
            run.bold = style["bold"]
            run.font.color.rgb = style["color"]
            run.font.name = "微软雅黑"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")

        elif tag == "p":
            if not element.get_text().strip():
                continue
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.line_spacing = 1.5
            _add_formatted_runs(p, element)

        elif tag in ("ul", "ol"):
            for i, li in enumerate(element.find_all("li", recursive=False)):
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Cm(1.5)
                p.paragraph_format.space_after = Pt(4)
                prefix = f"{i + 1}. " if tag == "ol" else "• "
                run = p.add_run(prefix)
                run.font.name = "微软雅黑"
                run.font.size = Pt(12)
                _add_formatted_runs(p, li)

        elif tag == "blockquote":
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(1)
            p.paragraph_format.space_after = Pt(6)
            run = p.add_run(element.get_text())
            run.italic = True
            run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
            run.font.name = "微软雅黑"
            run.font.size = Pt(12)
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")

        elif tag == "img":
            img_src = element.get("src", "")
            if img_src and not img_src.startswith(("http://", "https://")):
                img_path = Path(img_src)
                if img_path.exists():
                    doc.add_picture(str(img_path), width=Inches(5))

        elif tag == "table":
            rows = element.find_all("tr")
            if rows:
                header_cells = rows[0].find_all(["th", "td"])
                table = doc.add_table(rows=len(rows), cols=len(header_cells))
                table.style = "Light Grid Accent 1"
                for i, tr in enumerate(rows):
                    for j, cell in enumerate(tr.find_all(["th", "td"])):
                        table_cell = table.cell(i, j)
                        table_cell.text = cell.get_text().strip()
                        if i == 0:
                            for paragraph in table_cell.paragraphs:
                                for run in paragraph.runs:
                                    run.bold = True

        elif tag == "hr":
            doc.add_paragraph()

        elif tag == "br":
            pass

        elif tag == "strong" or tag == "b":
            if element.parent and element.parent.name not in ("p", "li", "td", "th"):
                p = doc.add_paragraph()
                run = p.add_run(element.get_text())
                run.bold = True
                run.font.name = "微软雅黑"
                run.font.size = Pt(12)

        elif tag == "em" or tag == "i":
            if element.parent and element.parent.name not in ("p", "li", "td", "th"):
                p = doc.add_paragraph()
                run = p.add_run(element.get_text())
                run.italic = True
                run.font.name = "微软雅黑"
                run.font.size = Pt(12)

        elif tag == "pre":
            code_text = element.get_text()
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            run = p.add_run(code_text)
            run.font.name = "Consolas"
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    output = Path(output_path)
    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output))


def _add_formatted_runs(paragraph, element):
    """Add formatted text runs to a Word paragraph from HTML element."""
    for child in element.children:
        if hasattr(child, "name"):
            if child.name == "strong" or child.name == "b":
                run = paragraph.add_run(child.get_text())
                run.bold = True
            elif child.name == "em" or child.name == "i":
                run = paragraph.add_run(child.get_text())
                run.italic = True
            elif child.name == "a":
                run = paragraph.add_run(child.get_text())
                run.font.color.rgb = RGBColor(0x1A, 0x73, 0xE8)
            elif child.name == "img":
                img_src = child.get("src", "")
                if img_src and not img_src.startswith(("http://", "https://")):
                    img_path = Path(img_src)
                    if img_path.exists():
                        paragraph.add_run()
                        try:
                            paragraph.runs[-1].add_picture(str(img_path), width=Inches(4))
                        except Exception:
                            pass
            elif child.name == "br":
                paragraph.add_run("\n")
            else:
                paragraph.add_run(child.get_text())
        else:
            text = str(child)
            if text.strip():
                paragraph.add_run(text)

    if not element.children:
        run = paragraph.add_run(element.get_text())
        run.font.name = "微软雅黑"
        run.font.size = Pt(12)
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")


def _generate_word_from_md(md_content: str, output_path: str):
    """Generate Word document directly from Markdown content."""
    import markdown
    html_body = markdown.markdown(
        md_content,
        extensions=["tables", "fenced_code", "codehilite", "toc"],
        output_format="html"
    )
    full_html = f"<html><body>{html_body}</body></html>"
    _generate_word_from_html(full_html, output_path)


def export_platform(md_content: str, platform: str, output_dir: str, theme_name: str = "professional-clean", images_dir: str = None):
    """Export content for a single platform.

    Args:
        md_content: Source Markdown content
        platform: Target platform name
        output_dir: Output directory path
        theme_name: Theme name for HTML conversion
        images_dir: Images directory path (copied to output)

    Returns:
        dict with exported file paths
    """
    output = Path(output_dir)
    output.mkdir(parents=True, exist_ok=True)

    files = {}
    platform_name = PLATFORM_DISPLAY_NAMES.get(platform, platform)
    base_name = f"{platform_name}-{platform}"

    processed_md = _process_markdown_for_platform(md_content, platform)

    md_path = output / f"{base_name}.md"
    md_path.write_text(processed_md, encoding="utf-8")
    files["md"] = str(md_path)

    if platform in PLATFORM_FORMATS and PLATFORM_FORMATS[platform] == "html":
        if platform == "wechat":
            html_content = _convert_to_html(processed_md, theme_name)
        else:
            html_content = _convert_to_standard_html(processed_md)
        html_path = output / f"{base_name}.html"
        html_path.write_text(html_content, encoding="utf-8")
        files["html"] = str(html_path)

    docx_path = output / f"{base_name}.docx"
    if platform in PLATFORM_FORMATS and PLATFORM_FORMATS[platform] == "html":
        html_content = files.get("html")
        if html_content:
            _generate_word_from_html(Path(html_content).read_text(encoding="utf-8"), str(docx_path))
        else:
            _generate_word_from_md(processed_md, str(docx_path))
    else:
        _generate_word_from_md(processed_md, str(docx_path))
    files["docx"] = str(docx_path)

    if images_dir:
        images_src = Path(images_dir)
        images_dst = output / "images"
        if images_src.exists():
            import shutil
            if images_src.resolve() == images_dst.resolve():
                files["images"] = str(images_dst)
            else:
                if images_dst.exists():
                    shutil.rmtree(images_dst)
                shutil.copytree(images_src, images_dst)
                files["images"] = str(images_dst)

    return files


def generate_usage_guide(output_dir: str, platform_files: dict, images_dir: str = None):
    """Generate a Markdown usage guide for exported files."""
    output = Path(output_dir)
    
    platform_instructions = {
        "wechat": {
            "desc": "微信公众号文章",
            "usage": "复制 HTML 内容到微信公众平台编辑器，图片需手动上传",
            "cover_ratio": "2.35:1",
        },
        "xiaohongshu": {
            "desc": "小红书笔记",
            "usage": "复制文字内容，配合 3:4 封面图发布",
            "cover_ratio": "3:4",
        },
        "zhihu": {
            "desc": "知乎文章/回答",
            "usage": "复制 Markdown 内容到知乎编辑器",
            "cover_ratio": "16:9",
        },
        "baijiahao": {
            "desc": "百度百家号文章",
            "usage": "复制 HTML 内容到百家号编辑器",
            "cover_ratio": "16:9",
        },
        "weibo": {
            "desc": "微博长文",
            "usage": "复制文字内容，配图发布（字数超限时使用头条文章）",
            "cover_ratio": "1:1",
        },
        "sohu": {
            "desc": "搜狐号文章",
            "usage": "复制 HTML 内容到搜狐号编辑器",
            "cover_ratio": "16:9",
        },
        "toutiao": {
            "desc": "今日头条文章",
            "usage": "复制 HTML 内容到头条号编辑器",
            "cover_ratio": "16:9",
        },
        "qiehao": {
            "desc": "企鹅号（腾讯新闻）文章",
            "usage": "复制 HTML 内容到企鹅号编辑器",
            "cover_ratio": "16:9",
        },
        "jianshu": {
            "desc": "简书文章",
            "usage": "复制 Markdown 内容到简书编辑器",
            "cover_ratio": "16:9",
        },
        "douban": {
            "desc": "豆瓣文章/日记",
            "usage": "复制文字内容到豆瓣编辑器",
            "cover_ratio": "16:9",
        },
        "dayu": {
            "desc": "大鱼号（UC）文章",
            "usage": "复制 HTML 内容到大鱼号编辑器",
            "cover_ratio": "16:9",
        },
        "kr36": {
            "desc": "36氪文章",
            "usage": "复制 Markdown 内容到36氪投稿系统",
            "cover_ratio": "16:9",
        },
        "bilibili": {
            "desc": "B站视频脚本/专栏",
            "usage": "按脚本录制视频，或作为专栏文章发布",
            "cover_ratio": "16:9",
        },
        "douyin": {
            "desc": "抖音短视频脚本",
            "usage": "按脚本拍摄短视频，配合文案发布",
            "cover_ratio": "9:16",
        },
        "newsletter": {
            "desc": "Newsletter 邮件",
            "usage": "发送 HTML 邮件，或复制文字内容到邮件平台",
            "cover_ratio": "N/A",
        },
    }
    
    format_map = {"md": "Markdown", "html": "HTML", "docx": "Word"}
    
    # Build MD content
    md_lines = [
        "# 📂 文件使用说明",
        "",
        "本文档说明本次导出的所有文件及其用途。",
        "",
        "**导航**：[📄 文件清单](文件清单_可点击预览.md) ｜ [🖼️ 图片清单](图片清单.md)",
        "",
        "## 📋 文件清单",
        "",
        "| 平台 | 文件 | 用途说明 |",
        "|------|------|----------|",
    ]
    
    for platform, files in platform_files.items():
        info = platform_instructions.get(platform, {"desc": platform, "usage": ""})
        file_links = []
        for fmt, fpath in files.items():
            if fmt == "images":
                continue
            fname = Path(fpath).name
            fmt_label = format_map.get(fmt, fmt.upper())
            file_links.append(f"`{fname}` ({fmt_label})")
        files_str = "、".join(file_links)
        md_lines.append(f"| {info['desc']} | {files_str} | {info['usage']} |")
    
    # Build image rows
    if images_dir:
        images_path = Path(images_dir)
        if images_path.exists():
            md_lines.append("")
            md_lines.append("## 🖼️ 配图文件")
            md_lines.append("")
            md_lines.append("| 文件名 | 说明 |")
            md_lines.append("|--------|------|")
            for img in sorted(images_path.iterdir()):
                if img.suffix.lower() in (".png", ".jpg", ".jpeg", ".webp"):
                    md_lines.append(f"| {img.name} | {img.name.replace('_', ' ')} |")
    
    md_lines.append("")
    md_lines.append("## 💡 发布提示")
    md_lines.append("")
    md_lines.append("1. **HTML 文件**：直接复制全部内容到对应平台的富文本编辑器")
    md_lines.append("2. **Markdown 文件**：复制文字内容到平台编辑器（部分平台支持 Markdown）")
    md_lines.append("3. **Word 文件**：可用 Word 打开查看，或用于需要上传文档的场景")
    md_lines.append("4. **图片**：需要手动上传到对应平台，然后插入文章中")
    md_lines.append("5. **视频脚本**：按脚本标注的分镜和时间拍摄视频")
    md_lines.append("")
    md_lines.append("---")
    md_lines.append("*由 Writer 自动生成*")
    md_lines.append("")
    
    readme_path = output / "README_文件说明.md"
    readme_path.write_text("\n".join(md_lines), encoding="utf-8")
    
    return str(readme_path)


def generate_image_gallery_html(output_dir: str, images_dir: str = None):
    """Generate a Markdown image gallery for preview."""
    output = Path(output_dir)
    if not images_dir:
        images_dir = output / "images"
    images_path = Path(images_dir)
    
    images = []
    if images_path.exists():
        for img in sorted(images_path.iterdir()):
            if img.suffix.lower() in (".png", ".jpg", ".jpeg", ".webp"):
                images.append({
                    "name": img.name,
                    "src": f"images/{img.name}",
                    "type": "封面" if "cover" in img.name else "配图",
                })
    
    md_lines = [
        "# 📸 图片清单",
        "",
        "**导航**：[📂 文件说明](README_文件说明.md) ｜ [📄 文件清单](文件清单_可点击预览.md)",
        "",
    ]
    
    if not images:
        md_lines.append("暂无图片文件。图片生成服务不可用，可手动生成后重新运行导出。")
    else:
        md_lines.append(f"## 全部图片（{len(images)} 张）")
        md_lines.append("")
        md_lines.append("| 文件名 | 类型 | 路径 |")
        md_lines.append("|--------|------|------|")
        for img in images:
            md_lines.append(f"| {img['name']} | {img['type']} | `{img['src']}` |")
    
    md_lines.append("")
    md_lines.append("---")
    md_lines.append("*由 Writer 自动生成*")
    md_lines.append("")
    
    gallery_path = output / "图片清单.md"
    gallery_path.write_text("\n".join(md_lines), encoding="utf-8")
    return str(gallery_path)


def generate_file_preview_html(output_dir: str, platform_files: dict, images_dir: str = None):
    """Generate a Markdown file list with preview info."""
    output = Path(output_dir)
    
    platform_info_map = {
        "wechat": {"name": "公众号", "desc": "微信公众号文章"},
        "xiaohongshu": {"name": "小红书", "desc": "小红书笔记"},
        "zhihu": {"name": "知乎", "desc": "知乎文章/回答"},
        "baijiahao": {"name": "百家号", "desc": "百度百家号文章"},
        "weibo": {"name": "微博", "desc": "微博长文"},
        "sohu": {"name": "搜狐", "desc": "搜狐号文章"},
        "toutiao": {"name": "今日头条", "desc": "今日头条文章"},
        "qiehao": {"name": "企鹅号", "desc": "企鹅号（腾讯新闻）文章"},
        "jianshu": {"name": "简书", "desc": "简书文章"},
        "douban": {"name": "豆瓣", "desc": "豆瓣文章/日记"},
        "dayu": {"name": "大鱼号", "desc": "大鱼号（UC）文章"},
        "kr36": {"name": "36氪", "desc": "36氪文章"},
        "bilibili": {"name": "B站", "desc": "B站视频脚本/专栏"},
        "douyin": {"name": "抖音", "desc": "抖音短视频脚本"},
        "newsletter": {"name": "Newsletter", "desc": "邮件订阅"},
    }
    
    format_map = {"md": "Markdown", "html": "HTML", "docx": "Word"}
    
    md_lines = [
        "# 📂 文件清单 - 可点击预览",
        "",
        f"输出目录: `{output.name}/`",
        "",
        "**导航**：[📂 文件说明](README_文件说明.md) ｜ [🖼️ 图片清单](图片清单.md)",
        "",
        "## 📄 快速导航",
        "",
        "| 文件 | 操作 |",
        "|------|------|",
        "| **README_文件说明.md** — 文件使用说明与发布指引 | [打开](README_文件说明.md) |",
        "| **图片清单.md** — 所有配图预览与下载 | [打开](图片清单.md) |",
        "",
    ]
    
    total_files = 0
    for platform, files in platform_files.items():
        info = platform_info_map.get(platform, {"name": platform, "desc": platform})
        md_lines.append(f"### {info['name']}（{info['desc']}）")
        md_lines.append("")
        md_lines.append("| 文件 | 格式 | 操作 |")
        md_lines.append("|------|------|------|")
        
        for fmt, fpath in files.items():
            if fmt == "images":
                continue
            fname = Path(fpath).name
            fmt_label = format_map.get(fmt, fmt.upper())
            total_files += 1
            if fmt == "docx":
                md_lines.append(f"| `{fname}` | {fmt_label} | [下载]({fname}) |")
            else:
                md_lines.append(f"| `{fname}` | {fmt_label} | [预览]({fname}) |")
        md_lines.append("")
    
    md_lines.append("## 🖼️ 图片文件")
    md_lines.append("")
    images_path = Path(images_dir) if images_dir else output / "images"
    image_count = 0
    if images_path.exists():
        md_lines.append("| 缩略图 | 文件名 | 类型 | 操作 |")
        md_lines.append("|--------|--------|------|------|")
        for img in sorted(images_path.iterdir()):
            if img.suffix.lower() in (".png", ".jpg", ".jpeg", ".webp"):
                is_cover = "cover" in img.name
                img_type = "封面" if is_cover else "配图"
                md_lines.append(f"| ![](images/{img.name}) | `{img.name}` | {img_type} | [查看原图](images/{img.name}) · [下载](images/{img.name}) |")
                image_count += 1
    
    if image_count == 0:
        md_lines.append("暂无图片文件。")
    
    md_lines.append("")
    md_lines.append("---")
    md_lines.append(f"*共 {len(platform_files)} 个平台、{total_files} 个文件、{image_count} 张图片*")
    md_lines.append("*由 Writer 自动生成*")
    md_lines.append("")
    
    preview_path = output / "文件清单_可点击预览.md"
    preview_path.write_text("\n".join(md_lines), encoding="utf-8")
    return str(preview_path)


def export_all_platforms(md_content: str, output_dir: str, theme_name: str = "professional-clean", images_dir: str = None):
    """Export content for all supported platforms.

    Args:
        md_content: Source Markdown content
        output_dir: Output directory path
        theme_name: Theme name for HTML conversion
        images_dir: Images directory path

    Returns:
        dict mapping platform -> exported file paths
    """
    results = {}
    for platform in SUPPORTED_PLATFORMS:
        results[platform] = export_platform(md_content, platform, output_dir, theme_name, images_dir)
    return results


if __name__ == "__main__":
    import argparse

    parser = argparse.ArgumentParser(description="Multi-platform file exporter")
    sub = parser.add_subparsers(dest="command", required=True)

    p_export = sub.add_parser("export", help="Export to platform-specific files")
    p_export.add_argument("input", help="Markdown file path")
    p_export.add_argument("--platform", required=True, help="Target platform or 'all'")
    p_export.add_argument("--output", required=True, help="Output directory")
    p_export.add_argument("--theme", default="professional-clean", help="Theme name for HTML platforms")
    p_export.add_argument("--images", help="Images directory to copy")

    args = parser.parse_args()

    md_content = Path(args.input).read_text(encoding="utf-8")
    style = _load_style()
    theme = args.theme or style.get("theme", "professional-clean")

    if args.platform == "all":
        results = export_all_platforms(md_content, args.output, theme, args.images)
        print(f"Exported {len(results)} platforms to {args.output}")
        for platform, files in results.items():
            print(f"  {platform}: {', '.join(files.values())}")
    else:
        if args.platform not in SUPPORTED_PLATFORMS:
            print(f"Error: unsupported platform '{args.platform}'", file=sys.stderr)
            print(f"Supported: {', '.join(SUPPORTED_PLATFORMS)}", file=sys.stderr)
            sys.exit(1)
        results = export_platform(md_content, args.platform, args.output, theme, args.images)
        print(f"Exported {args.platform} to {args.output}")
        for fmt, path in results.items():
            print(f"  {fmt}: {path}")
